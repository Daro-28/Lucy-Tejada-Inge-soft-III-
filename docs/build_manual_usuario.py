from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Manual_Usuario_Lucy_Tejada.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"
TEXT = RGBColor(31, 41, 55)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, sz="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_paragraph(doc, text="", style=None, bold_start=None):
    p = doc.add_paragraph(style=style)
    if bold_start and text.startswith(bold_start):
        add_run(p, bold_start, bold=True)
        add_run(p, text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Manual Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Manual Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_GRAY)
        hdr[i].paragraphs[0].add_run(header).bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].paragraphs[0].add_run(value)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Manual Bullet" not in styles:
        bullet = styles.add_style("Manual Bullet", 1)
    else:
        bullet = styles["Manual Bullet"]
    bullet.base_style = normal
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)

    if "Manual Number" not in styles:
        number = styles.add_style("Manual Number", 1)
    else:
        number = styles["Manual Number"]
    number.base_style = normal
    number.paragraph_format.left_indent = Inches(0.5)
    number.paragraph_format.first_line_indent = Inches(-0.25)
    number.paragraph_format.space_after = Pt(4)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Manual de Usuario - Sistema Centro Cultural Lucy Tejada")
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def add_cover(doc):
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Manual de Usuario\nSistema Centro Cultural Lucy Tejada")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Guía para estudiantes, educadores y administradores")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Proyecto de Ingeniería de Software III", bold=True, color=BLUE)
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    set_table_width(table, [2.0, 4.3])
    rows = [
        ("Aplicación", "Plataforma web para gestión de formación cultural"),
        ("Frontend", "React + TypeScript + Vite"),
        ("Backend", "NestJS + PostgreSQL + JWT"),
        ("Fecha del manual", date.today().strftime("%Y-%m-%d")),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = table.rows[idx].cells
        set_cell_shading(cells[0], LIGHT_BLUE)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(value)
    set_table_width(table, [2.0, 4.3])
    doc.add_page_break()


def build_manual():
    doc = Document()
    configure_styles(doc)

    add_cover(doc)

    doc.add_heading("1. Propósito del sistema", level=1)
    add_paragraph(
        doc,
        "El Sistema Centro Cultural Lucy Tejada permite consultar y administrar información relacionada con programas de formación artística, grupos, matrículas, asistencias, evaluaciones, escenarios, reservas, usuarios y notificaciones.",
    )
    add_paragraph(
        doc,
        "La aplicación está organizada por roles. Cada usuario ingresa con sus credenciales y el sistema lo redirige al panel que corresponde a su perfil: Estudiante, Educador o Administrador.",
    )

    add_table(
        doc,
        ["Rol", "Acceso principal", "Qué puede consultar"],
        [
            ("Estudiante", "/estudiante", "Cursos activos, asistencias, notas y notificaciones personales."),
            ("Educador", "/profesor", "Grupos asignados, estudiantes activos, asistencias, evaluaciones y alertas."),
            ("Administrador", "/admin", "Usuarios, programas, grupos, matrículas, escenarios, reservas y notificaciones."),
        ],
        [1.4, 1.6, 3.5],
    )

    doc.add_heading("2. Requisitos para usar el sistema", level=1)
    add_bullets(
        doc,
        [
            "Navegador web moderno.",
            "Backend ejecutándose en http://localhost:3000.",
            "Frontend ejecutándose en http://localhost:5173.",
            "Conexión a una base de datos PostgreSQL definida en backend/.env mediante DATABASE_URL.",
            "JWT_SECRET definido en backend/.env para firmar las sesiones.",
        ],
    )

    doc.add_heading("3. Inicio rápido del proyecto", level=1)
    add_paragraph(doc, "Estos pasos son para levantar la aplicación en ambiente local.")
    add_numbered(
        doc,
        [
            "Abrir una terminal en la carpeta backend y ejecutar npm install.",
            "Crear el archivo backend/.env con DATABASE_URL y JWT_SECRET.",
            "Ejecutar npm run seed para cargar usuarios y datos de prueba.",
            "Ejecutar npm run dev en backend para iniciar la API en el puerto 3000.",
            "Abrir otra terminal en la carpeta frontend y ejecutar npm install.",
            "Ejecutar npm run dev en frontend y abrir http://localhost:5173.",
        ],
    )

    add_table(
        doc,
        ["Servicio", "Comando", "Resultado esperado"],
        [
            ("Backend", "npm run dev", "API disponible en http://localhost:3000 y Swagger en /api."),
            ("Seed", "npm run seed", "Carga usuarios, programas, grupos, matrículas, asistencias y reservas de prueba."),
            ("Frontend", "npm run dev", "Interfaz disponible en http://localhost:5173."),
        ],
        [1.4, 1.7, 3.4],
    )

    doc.add_heading("4. Credenciales de prueba", level=1)
    add_paragraph(doc, "Después de ejecutar el seed, se pueden usar las siguientes cuentas:")
    add_table(
        doc,
        ["Perfil", "Correo", "Contraseña"],
        [
            ("Administrador", "admin@lucytejada.edu.co", "Prueba123"),
            ("Educador", "profesor@lucytejada.edu.co", "Prueba123"),
            ("Estudiante", "estudiante@lucytejada.edu.co", "Prueba123"),
        ],
        [1.6, 3.2, 1.7],
    )

    doc.add_heading("5. Pantalla de inicio", level=1)
    add_paragraph(
        doc,
        "La pantalla inicial presenta el Centro Cultural Lucy Tejada, los programas de formación y los accesos a Ingresar y Registrarse.",
    )
    add_bullets(
        doc,
        [
            "Inicio: muestra el mensaje principal del sistema.",
            "Programas: permite revisar las áreas de Música, Danza y Teatro, y Artes Plásticas.",
            "Ingresar: abre el formulario de inicio de sesión.",
            "Registrarse: abre el formulario para crear una cuenta nueva.",
            "Si el usuario ya inició sesión, aparece la opción Ir a mi panel.",
        ],
    )

    doc.add_heading("6. Registro de usuarios", level=1)
    add_paragraph(doc, "Para crear una cuenta nueva, entrar a Registrarse y completar el formulario.")
    add_numbered(
        doc,
        [
            "Escribir el nombre completo.",
            "Ingresar un correo electrónico válido.",
            "Seleccionar el rol: Estudiante, Educador / Profesor o Administrativo / Admin.",
            "Completar el campo adicional si aplica: código estudiantil para estudiante o especialidad artística para educador.",
            "Definir y confirmar una contraseña de mínimo 6 caracteres.",
            "Presionar Crear Cuenta. Si el registro es exitoso, el sistema redirige al login.",
        ],
    )

    doc.add_heading("7. Inicio de sesión", level=1)
    add_paragraph(doc, "Para entrar al sistema, abrir Ingresar y completar correo y contraseña.")
    add_bullets(
        doc,
        [
            "Si las credenciales son correctas, el sistema guarda la sesión y redirige automáticamente según el rol.",
            "Si el correo o la contraseña son incorrectos, se muestra un mensaje de error.",
            "Si el usuario está inactivo, el sistema no permite el acceso.",
            "La sesión usa token JWT y las rutas privadas verifican el rol antes de mostrar cada panel.",
        ],
    )

    doc.add_heading("8. Panel del Estudiante", level=1)
    add_paragraph(
        doc,
        "El panel del estudiante está pensado para consulta personal. Muestra un resumen de cursos inscritos, promedio de asistencia, nota promedio y notificaciones.",
    )
    add_table(
        doc,
        ["Sección", "Uso"],
        [
            ("Mis Matrículas", "Consultar cursos activos, programa, horario, docente asignado y estado de matrícula."),
            ("Mis Asistencias", "Revisar registros recientes de asistencia por grupo y fecha."),
            ("Mis Notas", "Consultar evaluaciones, descripción y calificación registrada."),
            ("Notificaciones", "Leer mensajes personales enviados por el sistema o administración."),
        ],
        [2.0, 4.5],
    )
    add_paragraph(doc, "Para salir del panel, usar el botón Cerrar Sesión ubicado en la barra lateral.")

    doc.add_heading("9. Panel del Educador", level=1)
    add_paragraph(
        doc,
        "El panel del educador resume los grupos asignados y la información académica relacionada con sus estudiantes.",
    )
    add_table(
        doc,
        ["Sección", "Uso"],
        [
            ("Mis Grupos", "Ver nombre del grupo, programa, horario, cupo máximo y estudiantes inscritos."),
            ("Asistencias", "Consultar las últimas asistencias registradas por estudiante, grupo y fecha."),
            ("Evaluaciones", "Revisar evaluaciones asociadas a los estudiantes de sus grupos."),
            ("Alertas", "Leer notificaciones personales del educador."),
        ],
        [2.0, 4.5],
    )
    add_paragraph(doc, "Los indicadores superiores muestran grupos asignados, estudiantes activos, asistencias registradas y evaluaciones.")

    doc.add_heading("10. Panel de Administración", level=1)
    add_paragraph(
        doc,
        "El panel administrativo ofrece una vista general del sistema y concentra la información de gestión institucional.",
    )
    add_table(
        doc,
        ["Sección", "Uso"],
        [
            ("Usuarios", "Consultar usuarios registrados y su rol."),
            ("Programas y Grupos", "Ver programas, grupos, horarios, cupos e instructor asignado."),
            ("Matrículas", "Consultar el volumen de matrículas activas desde los indicadores del panel."),
            ("Escenarios", "Revisar espacios, capacidad, reservas y estado: disponible, reservado o en mantenimiento."),
            ("Notificaciones", "Ver mensajes recientes asociados a usuarios o al sistema."),
        ],
        [2.0, 4.5],
    )
    add_paragraph(doc, "Los accesos CRUD del backend para administración están protegidos por JWT y roles, por lo que requieren iniciar sesión con una cuenta autorizada.")

    doc.add_heading("11. API y documentación técnica visible para usuarios avanzados", level=1)
    add_paragraph(
        doc,
        "El backend expone documentación Swagger en http://localhost:3000/api. Desde allí se pueden revisar rutas, probar endpoints y enviar el token Bearer obtenido en el login.",
    )
    add_table(
        doc,
        ["Módulo", "Rutas principales"],
        [
            ("Autenticación", "POST /auth/register, POST /auth/login"),
            ("Dashboard", "GET /dashboard/admin, /dashboard/profesor, /dashboard/estudiante"),
            ("Gestión", "usuarios, programas, grupos, estudiantes, educadores, matrículas"),
            ("Académico", "asistencias, evaluaciones, notificaciones"),
            ("Escenarios", "escenarios y reservas"),
        ],
        [1.6, 4.9],
    )

    doc.add_heading("12. Solución de problemas frecuentes", level=1)
    add_table(
        doc,
        ["Situación", "Qué revisar"],
        [
            ("No carga el panel", "Confirmar que el backend esté activo en el puerto 3000 y que el token de sesión sea válido."),
            ("Error de conexión", "Revisar VITE_API_URL en frontend o usar el valor por defecto http://localhost:3000."),
            ("Login fallido", "Verificar correo, contraseña, usuario activo y que JWT_SECRET exista en backend/.env."),
            ("No aparecen datos", "Ejecutar npm run seed y confirmar que DATABASE_URL apunte a la base correcta."),
            ("CORS", "El backend permite origen http://localhost:5173. Usar ese puerto o ajustar la configuración."),
        ],
        [2.0, 4.5],
    )

    doc.add_heading("13. Buenas prácticas de uso", level=1)
    add_bullets(
        doc,
        [
            "Cerrar sesión al terminar, especialmente en equipos compartidos.",
            "Usar correos reales o institucionales para identificar correctamente a cada usuario.",
            "Mantener actualizados los estados de matrícula, asistencia y evaluación desde los módulos administrativos o servicios correspondientes.",
            "Revisar Swagger solo si se tiene autorización para probar la API.",
            "No compartir contraseñas ni tokens de acceso.",
        ],
    )

    add_footer(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_manual()
